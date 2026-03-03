# ═══════════════════════════════════════════════════════════════════════════
#  ExamPro — Shobhit University Gangoh (SUG)
#  Team Believer © 2026
# ═══════════════════════════════════════════════════════════════════════════
import os
import uuid
import json
import re
import random
import hashlib
import datetime
import traceback
import io
from functools import wraps
from dotenv import load_dotenv
from flask import send_from_directory
from werkzeug.middleware.proxy_fix import ProxyFix

load_dotenv()

from flask import (
    Flask, render_template, request, session,
    jsonify, redirect, url_for
)
from flask_bcrypt import Bcrypt
from flask_socketio import SocketIO, emit, join_room, leave_room
from pymongo import MongoClient, ASCENDING, DESCENDING
from bson import ObjectId
from apscheduler.schedulers.background import BackgroundScheduler

# ── Optional: PyMuPDF ──────────────────────────────────────────────────────
try:
    import fitz
    PYMUPDF_OK = True
    print("[OK] PyMuPDF (fitz) ready")
except ImportError:
    PYMUPDF_OK = False
    print("[WARN] PyMuPDF not installed — trying pypdf")

# ── Optional: pypdf (Render fallback) ─────────────────────────────────────
try:
    from pypdf import PdfReader
    PYPDF_OK = True
    print("[OK] pypdf ready")
except ImportError:
    PYPDF_OK = False
    print("[WARN] pypdf not installed")

# ── Optional: python-docx ──────────────────────────────────────────────────
try:
    from docx import Document
    DOCX_OK = True
    print("[OK] python-docx ready")
except ImportError:
    DOCX_OK = False
    print("[WARN] python-docx not installed")

# ── Optional: Groq AI ──────────────────────────────────────────────────────
try:
    from groq import Groq as GroqClient
    GROQ_OK = True
except ImportError:
    GROQ_OK = False
    print("[WARN] groq not installed — pip install groq")

GROQ_KEY = os.environ.get('GROQ_API_KEY', '')
if GROQ_OK and GROQ_KEY:
    print(f"[OK] Groq AI ready — key: {GROQ_KEY[:8]}...{GROQ_KEY[-4:]}")
else:
    print("[WARN] GROQ_API_KEY not set — AI will use keyword fallback")

# ═══════════════════════════════════════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════════════════════════════════════
FACULTY_REG_KEY = os.environ.get('FACULTY_REGISTRATION_KEY', 'CS')
ROLL_REGEX      = re.compile(r'^\d{11}$')
ADMIN_USER      = os.environ.get('ADMIN_USERNAME',  'admin')
ADMIN_PASS      = os.environ.get('ADMIN_PASSWORD',  'Admin@SUG#2026')
IS_PRODUCTION   = os.environ.get('FLASK_ENV', '') == 'production'

# ═══════════════════════════════════════════════════════════════════════════
# APP SETUP
# ═══════════════════════════════════════════════════════════════════════════
app = Flask(__name__)

# ★ ProxyFix — Render/Azure reverse proxy ke liye ★
app.wsgi_app = ProxyFix(
    app.wsgi_app,
    x_for=1, x_proto=1, x_host=1, x_port=1
)

app.config['SECRET_KEY']                 = os.environ.get(
    'SECRET_KEY', 'exampro-sug-2026-xyz')
app.config['SESSION_COOKIE_NAME']        = 'exampro_sid'
app.config['SESSION_COOKIE_HTTPONLY']    = True
app.config['SESSION_COOKIE_SAMESITE']   = 'Lax'
# ★ Production mein HTTPS hai to Secure=True ★
app.config['SESSION_COOKIE_SECURE']     = IS_PRODUCTION
app.config['PERMANENT_SESSION_LIFETIME'] = datetime.timedelta(hours=8)

bcrypt   = Bcrypt(app)
socketio = SocketIO(
    app,
    cors_allowed_origins = '*',
    async_mode          = 'threading',
    logger              = False,
    engineio_logger     = False
)
print("[OK] App initialized")

# ═══════════════════════════════════════════════════════════════════════════
# MONGODB
# ═══════════════════════════════════════════════════════════════════════════
MONGO_URI = os.environ.get('MONGO_URI', '')
if not MONGO_URI:
    raise RuntimeError("MONGO_URI not set in .env!")

client = MongoClient(MONGO_URI, serverSelectionTimeoutMS=8000)
try:
    client.admin.command('ping')
    print("[OK] MongoDB connected")
except Exception as e:
    print(f"[ERROR] MongoDB: {e}")
    raise

db                = client['examdb']
users_col         = db['users']
exams_col         = db['exams']
questions_col     = db['questions']
answers_col       = db['answers']
marks_col         = db['marks']
results_col       = db['results']
violations_col    = db['violations']
notifications_col = db['notifications']

try:
    users_col.create_index(
        [('username', ASCENDING), ('role', ASCENDING)],
        unique=False, sparse=True)
    users_col.create_index(
        [('roll_number', ASCENDING), ('role', ASCENDING)],
        unique=False, sparse=True)
    users_col.create_index([('email', ASCENDING)], sparse=True)
    exams_col.create_index([('exam_id', ASCENDING)], unique=True)
    questions_col.create_index(
        [('exam_id', ASCENDING), ('index', ASCENDING)])
    answers_col.create_index(
        [('exam_id', ASCENDING), ('user_id', ASCENDING)])
    results_col.create_index(
        [('exam_id', ASCENDING), ('user_id', ASCENDING)])
    violations_col.create_index(
        [('exam_id', ASCENDING), ('user_id', ASCENDING)])
    print("[OK] Indexes ready")
except Exception:
    print("[INFO] Indexes already exist")

# ═══════════════════════════════════════════════════════════════════════════
# SCHEDULER
# ═══════════════════════════════════════════════════════════════════════════
scheduler = BackgroundScheduler(daemon=True)
scheduler.start()
print("[OK] Scheduler started")

# ═══════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════
def utcnow():
    return datetime.datetime.now(
        datetime.timezone.utc).replace(tzinfo=None)

def bson_to_dict(doc):
    if doc is None:
        return None
    doc = dict(doc)
    doc['_id'] = str(doc['_id'])
    for k, v in doc.items():
        if isinstance(v, datetime.datetime):
            doc[k] = v.isoformat()
        elif isinstance(v, ObjectId):
            doc[k] = str(v)
    return doc

def get_oid(s):
    try:
        return ObjectId(str(s))
    except Exception:
        return None

def device_fp():
    try:
        ua  = request.headers.get('User-Agent', '')
        ip  = request.remote_addr or ''
        raw = ua + ip + app.config['SECRET_KEY']
        return hashlib.sha256(raw.encode()).hexdigest()[:16]
    except Exception:
        return 'unknown'

def remaining_seconds(exam, ans_doc):
    try:
        duration = int(exam.get('duration_mins', 60)) * 60
        started  = ans_doc.get('start_time', utcnow())
        if isinstance(started, str):
            started = datetime.datetime.fromisoformat(started)
        elapsed = (utcnow() - started).total_seconds()
        return max(0, int(duration - elapsed))
    except Exception:
        return 0

# ═══════════════════════════════════════════════════════════════════════════
# AUTH DECORATORS
# ═══════════════════════════════════════════════════════════════════════════
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            if request.path.startswith('/api/'):
                return jsonify({'error': 'Unauthorized'}), 401
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated

def role_required(*roles):
    def decorator(f):
        @wraps(f)
        @login_required
        def decorated(*args, **kwargs):
            if session.get('role') not in roles:
                return jsonify({'error': 'Forbidden'}), 403
            return f(*args, **kwargs)
        return decorated
    return decorator

# ═══════════════════════════════════════════════════════════════════════════
# PAGE ROUTES
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/')
def index():
    if 'user_id' in session:
        return redirect(url_for('faculty_dashboard')
                        if session.get('role') == 'faculty'
                        else url_for('student_dashboard'))
    return redirect(url_for('login_page'))

@app.route('/login')
def login_page():
    return render_template('login.html')

@app.route('/register')
def register_page():
    return render_template('register.html')

@app.route('/student/dashboard')
@role_required('student')
def student_dashboard():
    return render_template('student_dashboard.html')

@app.route('/faculty/dashboard')
@role_required('faculty')
def faculty_dashboard():
    return render_template('faculty_dashboard.html')

@app.route('/exam/<exam_id>')
@role_required('student')
def exam_page(exam_id):
    return render_template('exam.html', exam_id=exam_id)

@app.route('/faculty/checking/<exam_id>')
@role_required('faculty')
def answer_checking(exam_id):
    return render_template('answer_checking.html',
                           exam_id=exam_id)

@app.route('/results')
@login_required
def results_page():
    return render_template('results.html')

@app.route('/results/<exam_id>')
@login_required
def results_exam_page(exam_id):
    return render_template('results.html', exam_id=exam_id)

@app.route('/ranking/<exam_id>')
@login_required
def ranking_page(exam_id):
    return render_template('ranking.html', exam_id=exam_id)

@app.route('/analytics/<exam_id>')
@role_required('faculty')
def analytics_page(exam_id):
    return render_template('analytics.html', exam_id=exam_id)

# ═══════════════════════════════════════════════════════════════════════════
# AUTH API
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/auth/register/student', methods=['POST'])
def register_student():
    try:
        data        = request.get_json(silent=True) or {}
        name        = data.get('name', '').strip()
        roll_number = data.get('roll_number', '').strip()
        email       = data.get('email', '').strip().lower()
        password    = data.get('password', '').strip()
        course      = data.get('course', '').strip()
        semester    = data.get('semester', '').strip()

        if not name:
            return jsonify(
                {'error': 'Full name is required'}), 400
        if not roll_number:
            return jsonify(
                {'error': 'Roll number is required'}), 400
        if not ROLL_REGEX.match(roll_number):
            return jsonify({'error':
                'Roll number must be exactly 11 digits '
                '(e.g. 23014168059)'}), 400
        if not password or len(password) < 6:
            return jsonify({'error':
                'Password must be at least 6 '
                'characters'}), 400
        if users_col.find_one(
                {'roll_number': roll_number,
                 'role': 'student'}):
            return jsonify({'error':
                f'Roll number {roll_number} is already '
                f'registered'}), 409
        if email and users_col.find_one({'email': email}):
            return jsonify(
                {'error': 'Email already registered'}), 409

        hashed = bcrypt.generate_password_hash(
            password).decode('utf-8')
        doc = {
            'user_id':     str(uuid.uuid4())[:8].upper(),
            'role':        'student',
            'name':        name,
            'username':    roll_number,
            'roll_number': roll_number,
            'email':       email,
            'password':    hashed,
            'course':      course,
            'semester':    semester,
            'is_active':   True,
            'created_at':  utcnow(),
            'last_login':  None
        }
        users_col.insert_one(doc)
        print(f"[REGISTER] ✅ Student: "
              f"{roll_number} — {name}")
        return jsonify({
            'message': f'Account created! Welcome {name}',
            'user_id': doc['user_id']
        }), 201
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/register/faculty', methods=['POST'])
def register_faculty():
    try:
        data       = request.get_json(silent=True) or {}
        secret_key = data.get('secret_key', '').strip()

        if secret_key != FACULTY_REG_KEY:
            print(f"[SECURITY] ❌ Wrong faculty key "
                  f"from {request.remote_addr}")
            return jsonify({'error':
                'Invalid registration key. '
                'Contact Admin / HOD.'}), 403

        name     = data.get('name', '').strip()
        username = data.get('username', '').strip().lower()
        email    = data.get('email', '').strip().lower()
        password = data.get('password', '').strip()
        emp_id   = data.get('employee_id', '').strip()

        if not name:
            return jsonify(
                {'error': 'Full name is required'}), 400
        if not username:
            return jsonify(
                {'error': 'Username is required'}), 400
        if not re.match(r'^[a-zA-Z0-9_.]+$', username):
            return jsonify({'error':
                'Username: only letters, numbers, '
                'dot, underscore'}), 400
        if not password or len(password) < 8:
            return jsonify({'error':
                'Faculty password must be at least '
                '8 characters'}), 400
        if users_col.find_one(
                {'username': username, 'role': 'faculty'}):
            return jsonify({'error':
                f'Username "{username}" already '
                f'taken'}), 409
        if email and users_col.find_one({'email': email}):
            return jsonify(
                {'error': 'Email already registered'}), 409

        hashed = bcrypt.generate_password_hash(
            password).decode('utf-8')
        doc = {
            'user_id':     str(uuid.uuid4())[:8].upper(),
            'role':        'faculty',
            'name':        name,
            'username':    username,
            'email':       email,
            'password':    hashed,
            'employee_id': emp_id,
            'department':  'Computer Science & Technology',
            'is_active':   True,
            'created_at':  utcnow(),
            'last_login':  None
        }
        users_col.insert_one(doc)
        print(f"[REGISTER] ✅ Faculty: "
              f"{username} — {name}")
        return jsonify({
            'message': f'Faculty account created! '
                       f'Welcome {name}',
            'user_id': doc['user_id']
        }), 201
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/register', methods=['POST'])
def api_register():
    data = request.get_json(silent=True) or {}
    role = data.get('role', 'student')
    if role == 'faculty':
        return register_faculty()
    return register_student()


@app.route('/api/auth/login', methods=['POST'])
def api_login():
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'Invalid JSON'}), 400

        role     = str(data.get('role',
                                'student')).strip()
        password = str(data.get('password', '')).strip()

        if not password:
            return jsonify(
                {'error': 'Password required'}), 400
        if role not in ['student', 'faculty']:
            return jsonify({'error': 'Invalid role'}), 400

        if role == 'student':
            roll = str(
                data.get('roll_number', '')).strip()
            if not roll:
                return jsonify(
                    {'error': 'Roll number required'}), 400
            if not ROLL_REGEX.match(roll):
                return jsonify({'error':
                    'Roll number must be exactly '
                    '11 digits'}), 400
            user = users_col.find_one(
                {'roll_number': roll, 'role': 'student'})
            if not user:
                return jsonify({'error':
                    'Student not found. Please '
                    'register first.'}), 404
        else:
            uname = str(
                data.get('username', '')).strip()
            if not uname:
                return jsonify(
                    {'error': 'Username required'}), 400
            user = users_col.find_one(
                {'username': uname, 'role': 'faculty'})
            if not user:
                return jsonify({'error':
                    'Faculty not found. Please '
                    'register first.'}), 404

        if not user.get('is_active', True):
            return jsonify({'error':
                'Account is deactivated. '
                'Contact Admin.'}), 403

        stored = user.get('password', '')
        if not stored:
            return jsonify({'error':
                'No password set for this '
                'account'}), 500
        if not bcrypt.check_password_hash(
                stored, password):
            return jsonify(
                {'error': 'Incorrect password'}), 401

        session.clear()
        session.permanent      = True
        session['user_id']     = str(user['_id'])
        session['role']        = user['role']
        session['username']    = str(
            user.get('username')
            or user.get('roll_number', ''))
        session['name']        = str(user.get('name', ''))
        session['fingerprint'] = device_fp()

        users_col.update_one(
            {'_id': user['_id']},
            {'$set': {'last_login': utcnow()}})

        redirect_url = (
            '/faculty/dashboard'
            if user['role'] == 'faculty'
            else '/student/dashboard'
        )
        print(f"[LOGIN] ✅ {user['role']}: "
              f"{session['username']} — {user['name']}")
        return jsonify({
            'message':  'Login successful',
            'role':     user['role'],
            'name':     session['name'],
            'redirect': redirect_url
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify(
            {'error': f'Server error: {str(e)}'}), 500


# ★ logout — GET + POST both support ★
@app.route('/api/auth/logout', methods=['POST', 'GET'])
def api_logout():
    name = session.get('name', 'User')
    session.clear()
    resp = jsonify({'message': f'Goodbye {name}!'})
    resp.delete_cookie(
        app.config['SESSION_COOKIE_NAME'],
        path='/'
    )
    return resp


@app.route('/api/auth/me')
@login_required
def api_me():
    return jsonify({
        'user_id':  session['user_id'],
        'role':     session['role'],
        'name':     session['name'],
        'username': session['username']
    })

# ═══════════════════════════════════════════════════════════════════════════
# EXAM API
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/exams', methods=['GET'])
@login_required
def get_exams():
    query = {}
    if session['role'] == 'student':
        query['status'] = {
            '$in': ['scheduled', 'live', 'ended']}
    exams = list(exams_col.find(
        query).sort('start_time', ASCENDING))
    return jsonify([bson_to_dict(e) for e in exams])


@app.route('/api/exams', methods=['POST'])
@role_required('faculty')
def create_exam():
    try:
        data    = request.get_json(silent=True) or {}
        exam_id = str(uuid.uuid4())[:8].upper()
        doc = {
            'exam_id':         exam_id,
            'title':           data.get('title', 'Untitled'),
            'subject':         data.get('subject', ''),
            'duration_mins':   int(
                data.get('duration_mins', 60)),
            'total_marks':     int(
                data.get('total_marks', 100)),
            'pass_marks':      int(
                data.get('pass_marks', 40)),
            'start_time':      datetime.datetime.fromisoformat(
                data['start_time']),
            'end_time':        datetime.datetime.fromisoformat(
                data['end_time']),
            'created_by':      session['user_id'],
            'status':          'scheduled',
            'randomize_q':     bool(
                data.get('randomize_q', True)),
            'randomize_o':     bool(
                data.get('randomize_o', True)),
            'violation_limit': int(
                data.get('violation_limit', 3)),
            'created_at':      utcnow()
        }
        exams_col.insert_one(doc)
        try:
            scheduler.add_job(
                auto_start_exam, 'date',
                run_date=doc['start_time'],
                args=[exam_id],
                id=f'start_{exam_id}',
                replace_existing=True)
            scheduler.add_job(
                auto_end_exam, 'date',
                run_date=doc['end_time'],
                args=[exam_id],
                id=f'end_{exam_id}',
                replace_existing=True)
        except Exception as e:
            print(f"[WARN] Scheduler: {e}")
        print(f"[EXAM] Created: {exam_id}")
        return jsonify({
            'exam_id': exam_id,
            'message': 'Exam created'
        }), 201
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/exams/<exam_id>', methods=['GET'])
@login_required
def get_exam(exam_id):
    exam = exams_col.find_one({'exam_id': exam_id})
    if not exam:
        return jsonify({'error': 'Exam not found'}), 404
    return jsonify(bson_to_dict(exam))


@app.route('/api/exams/<exam_id>/start', methods=['POST'])
@role_required('faculty')
def manual_start_exam(exam_id):
    auto_start_exam(exam_id)
    return jsonify({'message': f'Exam {exam_id} started'})


@app.route('/api/exams/<exam_id>/stop', methods=['POST'])
@role_required('faculty')
def stop_exam(exam_id):
    exam = exams_col.find_one({'exam_id': exam_id})
    if not exam:
        return jsonify({'error': 'Exam not found'}), 404
    exams_col.update_one(
        {'exam_id': exam_id},
        {'$set': {'status': 'ended'}})
    active = list(answers_col.find(
        {'exam_id': exam_id, 'submitted': False}))
    for ans in active:
        answers_col.update_one(
            {'_id': ans['_id']},
            {'$set': {
                'submitted':      True,
                'submit_time':    utcnow(),
                'auto_submitted': True
            }})
        auto_grade(exam_id, str(ans['user_id']))
    socketio.emit('exam_stopped', {
        'exam_id': exam_id,
        'title':   exam.get('title', 'Exam')
    })
    print(f"[EXAM] Stopped: {exam_id}, "
          f"auto-submitted {len(active)}")
    return jsonify({
        'message':        f'Exam {exam_id} stopped',
        'auto_submitted': len(active)
    })


@app.route('/api/exams/<exam_id>', methods=['DELETE'])
@role_required('faculty')
def delete_exam(exam_id):
    exam = exams_col.find_one({'exam_id': exam_id})
    if not exam:
        return jsonify({'error': 'Exam not found'}), 404
    title = exam.get('title', 'Exam')
    exams_col.delete_one({'exam_id': exam_id})
    questions_col.delete_many({'exam_id': exam_id})
    answers_col.delete_many({'exam_id': exam_id})
    marks_col.delete_many({'exam_id': exam_id})
    results_col.delete_many({'exam_id': exam_id})
    violations_col.delete_many({'exam_id': exam_id})
    try:
        scheduler.remove_job(f'start_{exam_id}')
    except Exception:
        pass
    try:
        scheduler.remove_job(f'end_{exam_id}')
    except Exception:
        pass
    socketio.emit('exam_deleted',
                  {'exam_id': exam_id, 'title': title})
    print(f"[EXAM] Deleted: {exam_id} — {title}")
    return jsonify(
        {'message': f'Exam "{title}" deleted successfully'})


def auto_start_exam(exam_id):
    exam = exams_col.find_one({'exam_id': exam_id})
    exams_col.update_one(
        {'exam_id': exam_id},
        {'$set': {'status': 'live'}})
    socketio.emit('exam_started', {
        'exam_id': exam_id,
        'title':   exam.get('title', '') if exam else ''
    })
    print(f"[EXAM] Started: {exam_id}")


def auto_end_exam(exam_id):
    exams_col.update_one(
        {'exam_id': exam_id},
        {'$set': {'status': 'ended'}})
    active = list(answers_col.find(
        {'exam_id': exam_id, 'submitted': False}))
    for ans in active:
        answers_col.update_one(
            {'_id': ans['_id']},
            {'$set': {
                'submitted':      True,
                'submit_time':    utcnow(),
                'auto_submitted': True
            }})
        auto_grade(exam_id, str(ans['user_id']))
    socketio.emit('exam_ended', {'exam_id': exam_id})
    print(f"[EXAM] Ended: {exam_id}, "
          f"auto-submitted {len(active)}")

# ═══════════════════════════════════════════════════════════════════════════
# AI CORE
# ═══════════════════════════════════════════════════════════════════════════
def get_groq_client():
    if GROQ_OK and GROQ_KEY:
        return GroqClient(api_key=GROQ_KEY)
    return None


def ai_parse_questions(text, subject='General'):
    client_ai = get_groq_client()
    if not client_ai:
        print("[AI Parse] Groq not available")
        return None

    prompt = f"""You are an expert university exam question extractor and answer writer.

Subject: {subject}

Your job:
1. Extract ALL questions from the text
2. For EVERY question, write a detailed model answer
3. Identify MCQ vs Subjective automatically
4. Estimate marks based on question complexity

Return ONLY a valid JSON array. No explanation. No markdown. No code blocks.

JSON Format:
[
  {{
    "text": "Complete question text here",
    "type": "mcq",
    "marks": 1,
    "options": ["Option A text", "Option B text", "Option C text", "Option D text"],
    "correct_option": "Option B text",
    "model_answer": "The correct answer is Option B because..."
  }},
  {{
    "text": "Subjective question text here",
    "type": "subjective",
    "marks": 5,
    "options": [],
    "correct_option": "",
    "model_answer": "Complete detailed answer..."
  }}
]

Rules:
- type MUST be "mcq" if options exist
- type MUST be "subjective" for descriptive questions
- options = texts WITHOUT a. b. c. d. prefix
- correct_option = exact text of correct answer
- model_answer = ALWAYS write detailed answer
- marks: 1 MCQ, 2-5 short, 5-10 long
- Do NOT skip any question

Text:
\"\"\"{text[:8000]}\"\"\"

JSON array only:"""

    try:
        resp = client_ai.chat.completions.create(
            model       = 'llama3-70b-8192',
            messages    = [{'role':    'user',
                            'content': prompt}],
            temperature = 0.15,
            max_tokens  = 4096
        )
        raw = resp.choices[0].message.content.strip()
        print(f"[AI Parse] Response: {len(raw)} chars")
        raw = re.sub(r'```json\s*', '', raw)
        raw = re.sub(r'```\s*',     '', raw)
        raw = raw.strip()
        start = raw.find('[')
        end   = raw.rfind(']') + 1
        if start == -1 or end <= start:
            print("[AI Parse] No JSON array")
            return None
        questions = json.loads(raw[start:end])
        cleaned   = []
        for q in questions:
            txt = str(q.get('text', '')).strip()
            if not txt or len(txt) < 5:
                continue
            qtype = str(
                q.get('type', 'subjective')).lower()
            if qtype not in ('mcq', 'subjective'):
                qtype = 'subjective'
            opts = [str(o).strip()
                    for o in q.get('options', [])
                    if str(o).strip()]
            if len(opts) >= 2:
                qtype = 'mcq'
            cleaned.append({
                'text':           txt,
                'type':           qtype,
                'marks':          max(1, int(
                    q.get('marks', 1))),
                'options':        opts,
                'correct_option': str(
                    q.get('correct_option',
                          '')).strip(),
                'model_answer':   str(
                    q.get('model_answer',
                          '')).strip()
            })
        print(f"[AI Parse] ✅ {len(cleaned)} questions")
        return cleaned if cleaned else None
    except json.JSONDecodeError as e:
        print(f"[AI Parse] JSON error: {e}")
        return None
    except Exception as e:
        print(f"[AI Parse] Error: {e}")
        traceback.print_exc()
        return None


def ai_generate_model_answers(questions,
                               subject='General'):
    client_ai = get_groq_client()
    if not client_ai:
        return questions
    need_answers = [
        (i, q) for i, q in enumerate(questions)
        if not q.get('model_answer', '').strip()
    ]
    if not need_answers:
        return questions
    print(f"[AI Answers] Generating for "
          f"{len(need_answers)} questions...")
    batch_size = 5
    for batch_start in range(
            0, len(need_answers), batch_size):
        batch = need_answers[
            batch_start:batch_start + batch_size]
        q_text = '\n\n'.join([
            f"Q{idx+1} ({q['type']}, "
            f"{q['marks']} marks): {q['text']}"
            + (f"\nOptions: {', '.join(q['options'])}"
               if q.get('options') else '')
            for idx, (_, q) in enumerate(batch)
        ])
        prompt = (
            f"You are an expert {subject} professor.\n"
            f"Write detailed model answers.\n\n"
            f"{q_text}\n\n"
            f"Return ONLY JSON array:\n"
            f'[{{"answer": "..."}}]'
        )
        try:
            resp = client_ai.chat.completions.create(
                model       = 'llama3-70b-8192',
                messages    = [{'role':    'user',
                                'content': prompt}],
                temperature = 0.2,
                max_tokens  = 3000
            )
            raw = resp.choices[0].message.content.strip()
            raw = re.sub(r'```json\s*', '', raw)
            raw = re.sub(r'```\s*',     '', raw)
            start   = raw.find('[')
            end     = raw.rfind(']') + 1
            answers = json.loads(raw[start:end])
            for j, (orig_idx, _) in enumerate(batch):
                if j < len(answers):
                    ans = str(
                        answers[j].get(
                            'answer', '')).strip()
                    if ans:
                        questions[orig_idx][
                            'model_answer'] = ans
        except Exception as e:
            print(f"[AI Answers] Batch error: {e}")
            continue
    return questions


def fallback_parse_questions(text):
    lines   = text.split('\n')
    result  = []
    cur     = None
    opt_buf = []
    q_re    = re.compile(
        r'^(?:Q\.?\s*)?(\d+)[.)]\s+(.+)',
        re.IGNORECASE)
    opt_re  = re.compile(
        r'^[a-d][.)]\s+', re.IGNORECASE)
    for line in lines:
        line = line.strip()
        if not line or len(line) < 4:
            continue
        m = q_re.match(line)
        if m:
            if cur:
                if len(opt_buf) >= 2:
                    cur['options'] = [
                        re.sub(r'^[a-d][.)]\s+', '',
                               o, flags=re.I).strip()
                        for o in opt_buf
                    ]
                    cur['type'] = 'mcq'
                result.append(cur)
                opt_buf = []
            cur = {
                'text':           m.group(2).strip(),
                'type':           'subjective',
                'marks':          1,
                'options':        [],
                'correct_option': '',
                'model_answer':   ''
            }
        elif opt_re.match(line) and cur:
            opt_buf.append(line)
        elif cur and len(line) > 10:
            cur['text'] += ' ' + line
    if cur:
        if len(opt_buf) >= 2:
            cur['options'] = [
                re.sub(r'^[a-d][.)]\s+', '',
                       o, flags=re.I).strip()
                for o in opt_buf
            ]
            cur['type'] = 'mcq'
        result.append(cur)
    if not result:
        result = [
            {
                'text':           l.strip(),
                'type':           'subjective',
                'marks':          1,
                'options':        [],
                'correct_option': '',
                'model_answer':   ''
            }
            for l in lines
            if l.strip() and len(l.strip()) > 10
        ]
    return result

# ═══════════════════════════════════════════════════════════════════════════
# QUESTIONS API
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/exams/<exam_id>/questions',
           methods=['POST'])
@role_required('faculty')
def save_questions(exam_id):
    try:
        data = request.get_json(silent=True) or {}
        qs   = data.get('questions', [])
        questions_col.delete_many({'exam_id': exam_id})
        docs = [{
            'exam_id':        exam_id,
            'index':          i,
            'type':           q.get('type', 'subjective'),
            'text':           q.get('text', ''),
            'options':        q.get('options', []),
            'correct_option': q.get(
                'correct_option', None),
            'model_answer':   q.get('model_answer', ''),
            'marks':          int(q.get('marks', 1)),
            'created_at':     utcnow()
        } for i, q in enumerate(qs)]
        if docs:
            questions_col.insert_many(docs)
        return jsonify(
            {'message': f'{len(docs)} questions saved'})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/exams/<exam_id>/questions/upload',
           methods=['POST'])
@role_required('faculty')
def upload_file(exam_id):
    """
    ★ FILE UPLOAD — PDF/DOCX/TXT ★
    PDF   → pypdf (Render) or fitz (local)
    DOCX  → python-docx
    TXT   → plain text
    """
    f = request.files.get('file')
    if not f:
        return jsonify({'error': 'No file uploaded'}), 400

    fname = (f.filename or '').lower().strip()
    text  = ''

    try:
        # ══════════════════════════════════════════
        # PDF
        # ══════════════════════════════════════════
        if fname.endswith('.pdf'):
            raw_bytes = f.read()

            if PYMUPDF_OK:
                # Local: fitz use karo
                doc = fitz.open(
                    stream=raw_bytes, filetype='pdf')
                for page in doc:
                    text += page.get_text('text') + '\n'
                doc.close()
                print(f"[PDF] fitz: "
                      f"{len(text)} chars")

            elif PYPDF_OK:
                # Render: pypdf use karo
                reader = PdfReader(io.BytesIO(raw_bytes))
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + '\n'
                print(f"[PDF] pypdf: "
                      f"{len(reader.pages)} pages, "
                      f"{len(text)} chars")

            else:
                return jsonify({'error':
                    'PDF library not available. '
                    'Please upload a .txt or '
                    '.docx file instead.'}), 500

        # ══════════════════════════════════════════
        # DOCX
        # ══════════════════════════════════════════
        elif fname.endswith('.docx'):
            if not DOCX_OK:
                return jsonify({'error':
                    'DOCX library not available. '
                    'Please upload a .txt '
                    'file instead.'}), 500
            raw_bytes = f.read()
            doc       = Document(io.BytesIO(raw_bytes))
            lines     = []
            for p in doc.paragraphs:
                if p.text.strip():
                    lines.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(
                                cell.text.strip())
            text = '\n'.join(lines)
            print(f"[DOCX] {len(lines)} lines, "
                  f"{len(text)} chars")

        # ══════════════════════════════════════════
        # TXT
        # ══════════════════════════════════════════
        elif fname.endswith('.txt'):
            text = f.read().decode('utf-8',
                                   errors='ignore')
            print(f"[TXT] {len(text)} chars")

        else:
            return jsonify({'error':
                'Only PDF, DOCX, TXT '
                'supported'}), 400

    except Exception as e:
        traceback.print_exc()
        return jsonify(
            {'error': f'File read error: {str(e)}'}), 500

    text = text.strip()
    if not text or len(text) < 15:
        return jsonify({'error':
            'No readable text found in file. '
            'Try a .txt file with questions.'}), 400

    print(f"[UPLOAD] '{f.filename}' → "
          f"{len(text)} chars | exam={exam_id}")

    exam    = exams_col.find_one({'exam_id': exam_id})
    subject = ((exam.get('subject') or 'General')
               if exam else 'General')

    # ── AI Parse ──────────────────────────────────
    questions = None
    try:
        questions = ai_parse_questions(text, subject)
    except Exception as e:
        print(f"[AI Parse] Exception: {e}")
        questions = None

    if questions:
        empty_ans = sum(
            1 for q in questions
            if not q.get('model_answer', '').strip()
        )
        if empty_ans > 0:
            try:
                questions = ai_generate_model_answers(
                    questions, subject)
            except Exception as e:
                print(f"[AI Answers] Error: {e}")

        print(f"[UPLOAD] ✅ AI: "
              f"{len(questions)} questions")
        return jsonify({
            'questions': questions,
            'count':     len(questions),
            'ai':        True,
            'subject':   subject,
            'message':   (
                f'AI extracted {len(questions)} '
                f'questions with model answers ✅'
            )
        })

    # ── Fallback ────────────────��─────────────────
    print("[UPLOAD] AI failed — fallback parser")
    questions = fallback_parse_questions(text)
    return jsonify({
        'questions': questions,
        'count':     len(questions),
        'ai':        False,
        'message':   (
            f'Extracted {len(questions)} questions '
            f'(AI unavailable — check GROQ_API_KEY)'
        )
    })


@app.route('/api/exams/<exam_id>/questions',
           methods=['GET'])
@login_required
def get_questions(exam_id):
    exam = exams_col.find_one({'exam_id': exam_id})
    if not exam:
        return jsonify({'error': 'Exam not found'}), 404
    qs = list(questions_col.find(
        {'exam_id': exam_id}
    ).sort('index', ASCENDING))
    if session['role'] == 'student':
        try:
            seed = (
                int(session['user_id'][:8], 16)
                ^ (hash(exam_id) & 0xFFFFFFFF)
            )
        except Exception:
            seed = hash(session['user_id'] + exam_id)
        rng = random.Random(seed)
        if exam.get('randomize_q'):
            rng.shuffle(qs)
        result = []
        for q in qs:
            q = bson_to_dict(q)
            q.pop('correct_option', None)
            q.pop('model_answer',   None)
            if (exam.get('randomize_o')
                    and q.get('options')):
                opts = q['options'][:]
                rng.shuffle(opts)
                q['options'] = opts
            result.append(q)
        return jsonify(result)
    return jsonify([bson_to_dict(q) for q in qs])

# ═══════════════════════════════════════════════════════════════════════════
# STUDENT SESSION / EXAM
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/exams/<exam_id>/session',
           methods=['POST'])
@role_required('student')
def start_session(exam_id):
    exam = exams_col.find_one({'exam_id': exam_id})
    if not exam:
        return jsonify({'error': 'Exam not found'}), 404
    if exam['status'] not in ('live', 'scheduled'):
        return jsonify(
            {'error': 'Exam not active'}), 400
    uid      = session['user_id']
    existing = answers_col.find_one(
        {'exam_id': exam_id, 'user_id': uid})
    if existing and existing.get('submitted'):
        return jsonify({
            'error':     'Already submitted',
            'submitted': True
        }), 409
    if not existing:
        doc = {
            'exam_id':        exam_id,
            'user_id':        uid,
            'answers':        {},
            'submitted':      False,
            'start_time':     utcnow(),
            'last_seen':      utcnow(),
            'ip':             request.remote_addr,
            'fingerprint':    device_fp(),
            'auto_submitted': False
        }
        answers_col.insert_one(doc)
        existing = doc
    return jsonify({
        'message':   'Session ready',
        'answers':   existing.get('answers', {}),
        'remaining': remaining_seconds(exam, existing),
        'submitted': existing.get('submitted', False)
    })


@app.route('/api/exams/<exam_id>/autosave',
           methods=['POST'])
@role_required('student')
def autosave(exam_id):
    data = request.get_json(silent=True) or {}
    uid  = session['user_id']
    answers_col.update_one(
        {'exam_id': exam_id, 'user_id': uid},
        {'$set': {
            'answers':   data.get('answers', {}),
            'last_seen': utcnow()
        }})
    return jsonify({'message': 'Saved'})


@app.route('/api/exams/<exam_id>/submit',
           methods=['POST'])
@role_required('student')
def submit_exam(exam_id):
    uid = session['user_id']
    rec = answers_col.find_one(
        {'exam_id': exam_id, 'user_id': uid})
    if not rec:
        return jsonify(
            {'error': 'No session found'}), 404
    if rec.get('submitted'):
        return jsonify(
            {'error': 'Already submitted'}), 409
    data = request.get_json(silent=True) or {}
    answers_col.update_one(
        {'exam_id': exam_id, 'user_id': uid},
        {'$set': {
            'answers':     data.get(
                'answers', rec.get('answers', {})),
            'submitted':   True,
            'submit_time': utcnow()
        }})
    auto_grade(exam_id, uid)
    socketio.emit('submission_alert', {
        'exam_id':  exam_id,
        'user_id':  uid,
        'username': session['username']
    })
    return jsonify({'message': 'Submitted successfully'})


@app.route('/api/exams/<exam_id>/heartbeat',
           methods=['POST'])
@role_required('student')
def heartbeat(exam_id):
    uid = session['user_id']
    answers_col.update_one(
        {'exam_id': exam_id, 'user_id': uid},
        {'$set': {'last_seen': utcnow()}})
    exam = exams_col.find_one({'exam_id': exam_id})
    rec  = answers_col.find_one(
        {'exam_id': exam_id, 'user_id': uid})
    if not exam or not rec:
        return jsonify({'error': 'Not found'}), 404
    return jsonify({
        'remaining': remaining_seconds(exam, rec),
        'status':    exam['status']
    })


def auto_grade(exam_id, user_id):
    rec      = answers_col.find_one(
        {'exam_id': exam_id, 'user_id': user_id})
    qs       = list(questions_col.find(
        {'exam_id': exam_id}))
    ans      = rec.get('answers', {}) if rec else {}
    total    = 0
    obtained = 0
    for q in qs:
        m      = int(q.get('marks', 1))
        total += m
        if (q['type'] == 'mcq'
                and str(ans.get(str(q['_id']),
                                '')).strip()
                == str(q.get('correct_option',
                             '')).strip()):
            obtained += m
    marks_col.update_one(
        {'exam_id': exam_id, 'user_id': user_id},
        {'$set': {
            'exam_id':      exam_id,
            'user_id':      user_id,
            'auto_marks':   obtained,
            'manual_marks': 0,
            'total_marks':  total,
            'remarks':      '',
            'checked_by':   None,
            'published':    False,
            'created_at':   utcnow()
        }},
        upsert=True
    )
    print(f"[AutoGrade] exam={exam_id} "
          f"user={user_id[:8]} "
          f"obtained={obtained}/{total}")
    return obtained

# ═══════════════════════════════════════════════════════════════════════════
# FACULTY MARKING
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/exams/<exam_id>/submissions',
           methods=['GET'])
@role_required('faculty')
def get_submissions(exam_id):
    subs = list(answers_col.find(
        {'exam_id': exam_id, 'submitted': True}))
    out  = []
    for s in subs:
        s   = bson_to_dict(s)
        oid = get_oid(s['user_id'])
        u   = (users_col.find_one({'_id': oid})
               if oid else None)
        s['student_name'] = (
            u.get('name', 'N/A') if u else 'N/A')
        s['roll_number']  = (
            u.get('roll_number', '') if u else '')
        out.append(s)
    return jsonify(out)


@app.route('/api/exams/<exam_id>/marks/<user_id>',
           methods=['GET'])
@role_required('faculty')
def get_marks(exam_id, user_id):
    m = marks_col.find_one(
        {'exam_id': exam_id, 'user_id': user_id})
    return jsonify(bson_to_dict(m) if m else {})


@app.route('/api/exams/<exam_id>/marks/<user_id>',
           methods=['PUT'])
@role_required('faculty')
def update_marks(exam_id, user_id):
    data = request.get_json(silent=True) or {}
    marks_col.update_one(
        {'exam_id': exam_id, 'user_id': user_id},
        {'$set': {
            'manual_marks': int(
                data.get('manual_marks', 0)),
            'remarks':      data.get('remarks', ''),
            'checked_by':   session['user_id'],
            'checked_at':   utcnow()
        }},
        upsert=True
    )
    return jsonify({'message': 'Marks saved'})


@app.route('/api/exams/<exam_id>/publish',
           methods=['POST'])
@role_required('faculty')
def publish_results(exam_id):
    exam = exams_col.find_one({'exam_id': exam_id})
    subs = list(answers_col.find(
        {'exam_id': exam_id, 'submitted': True}))
    rows = []
    for s in subs:
        uid = s['user_id']
        m   = marks_col.find_one(
            {'exam_id': exam_id, 'user_id': uid})
        if not m:
            continue
        total = (m.get('auto_marks', 0)
                 + m.get('manual_marks', 0))
        rows.append({'user_id': uid, 'total': total})
    rows.sort(key=lambda x: x['total'], reverse=True)
    n = len(rows)
    for i, r in enumerate(rows):
        rank       = i + 1
        percentile = (
            round((1 - (rank - 1) / n) * 100, 2)
            if n else 0
        )
        results_col.update_one(
            {'exam_id': exam_id,
             'user_id': r['user_id']},
            {'$set': {
                'exam_id':      exam_id,
                'user_id':      r['user_id'],
                'total':        r['total'],
                'rank':         rank,
                'percentile':   percentile,
                'pass':         r['total'] >= (
                    exam.get('pass_marks', 0)
                    if exam else 0),
                'published':    True,
                'published_at': utcnow()
            }},
            upsert=True
        )
        marks_col.update_one(
            {'exam_id': exam_id,
             'user_id': r['user_id']},
            {'$set': {'published': True}}
        )
    socketio.emit('result_published',
                  {'exam_id': exam_id})
    return jsonify(
        {'message':
         f'Results published for {n} students'})

# ═══════════════════════════════════════════════════════════════════════════
# RESULTS & RANKING
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/results', methods=['GET'])
@login_required
def get_results():
    if session['role'] == 'student':
        docs = list(results_col.find(
            {'user_id':   session['user_id'],
             'published': True}))
    else:
        docs = list(results_col.find(
            {'published': True}))
    return jsonify([bson_to_dict(d) for d in docs])


@app.route('/api/results/<exam_id>', methods=['GET'])
@login_required
def get_exam_results(exam_id):
    docs = list(results_col.find(
        {'exam_id': exam_id, 'published': True}
    ).sort('rank', ASCENDING))
    out = []
    for d in docs:
        d   = bson_to_dict(d)
        oid = get_oid(d['user_id'])
        u   = (users_col.find_one({'_id': oid})
               if oid else None)
        d['name']        = (
            u.get('name', '') if u else '')
        d['roll_number'] = (
            u.get('roll_number', '') if u else '')
        out.append(d)
    return jsonify(out)

# ══════════════════���════════════════════════════════════════════════════════
# VIOLATIONS
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/exams/<exam_id>/violation',
           methods=['POST'])
@role_required('student')
def log_violation(exam_id):
    data  = request.get_json(silent=True) or {}
    uid   = session['user_id']
    vtype = data.get('type', 'unknown')
    violations_col.insert_one({
        'exam_id':     exam_id,
        'user_id':     uid,
        'type':        vtype,
        'timestamp':   utcnow(),
        'ip':          request.remote_addr,
        'fingerprint': device_fp()
    })
    count = violations_col.count_documents(
        {'exam_id': exam_id, 'user_id': uid})
    exam  = exams_col.find_one({'exam_id': exam_id})
    limit = (exam.get('violation_limit', 3)
             if exam else 3)
    socketio.emit('violation_alert', {
        'exam_id':  exam_id,
        'user_id':  uid,
        'username': session['username'],
        'type':     vtype,
        'count':    count
    })
    if count >= limit:
        answers_col.update_one(
            {'exam_id': exam_id, 'user_id': uid},
            {'$set': {
                'submitted':      True,
                'submit_time':    utcnow(),
                'auto_submitted': True
            }})
        auto_grade(exam_id, uid)
        return jsonify({
            'auto_submit': True,
            'message':     'Violation limit reached'
        })
    return jsonify({'count': count, 'limit': limit})


@app.route('/api/exams/<exam_id>/violations',
           methods=['GET'])
@role_required('faculty')
def get_violations(exam_id):
    docs = list(violations_col.find(
        {'exam_id': exam_id}
    ).sort('timestamp', DESCENDING))
    return jsonify([bson_to_dict(d) for d in docs])

# ═══════════════════════════════════════════════════════════════════════════
# ANALYTICS
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/analytics/<exam_id>', methods=['GET'])
@role_required('faculty')
def analytics(exam_id):
    exam   = exams_col.find_one({'exam_id': exam_id})
    rdocs  = list(results_col.find(
        {'exam_id': exam_id}))
    scores = [r['total'] for r in rdocs]
    n      = len(scores)
    if n == 0:
        return jsonify(
            {'error': 'No results yet'}), 404
    avg       = round(sum(scores) / n, 2)
    passed    = sum(
        1 for r in rdocs if r.get('pass'))
    pass_rate = round(passed / n * 100, 2)
    pm        = (exam.get('total_marks', 100)
                 if exam else 100)
    buckets   = [0] * 10
    for s in scores:
        buckets[min(int(s / pm * 10), 9)] += 1
    viols = list(violations_col.aggregate([
        {'$match': {'exam_id': exam_id}},
        {'$group': {
            '_id':   '$type',
            'count': {'$sum': 1}
        }}
    ]))
    return jsonify({
        'total_students':      n,
        'average_score':       avg,
        'pass_rate':           pass_rate,
        'min_score':           min(scores),
        'max_score':           max(scores),
        'score_distribution':  buckets,
        'violation_analytics': [
            {'type': v['_id'], 'count': v['count']}
            for v in viols
        ]
    })

# ═══════════════════════════════════════════════════════════════════════════
# NOTIFICATIONS
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/api/notifications', methods=['GET'])
@login_required
def get_notifications():
    docs = list(notifications_col.find(
        {'$or': [
            {'to': session['user_id']},
            {'to': 'all'}
        ]}
    ).sort('created_at', DESCENDING).limit(20))
    return jsonify([bson_to_dict(d) for d in docs])

# ═══════════════════════════════════════════════════════════════════════════
# AI GRADING
# ═══════════════════════════════════════════════════════════════════════════
def ai_grade(question, answer, marks,
             subject='General'):
    if not answer or not str(answer).strip():
        return {
            'score':      0,
            'out_of':     marks,
            'feedback':   'No answer was provided.',
            'confidence': 'high',
            'key_points': ['Student did not answer'],
            'model':      'none'
        }
    client_ai = get_groq_client()
    if client_ai:
        try:
            prompt = (
                f"You are a strict but fair examiner "
                f"for {subject}.\n\n"
                f"QUESTION ({marks} marks): {question}\n"
                f"STUDENT ANSWER: {answer}\n\n"
                f"Respond ONLY with this JSON:\n"
                f'{{"score": <0-{marks}>, '
                f'"feedback": "<2-3 sentences>", '
                f'"confidence": "<high|medium|low>", '
                f'"key_points": ["<p1>", "<p2>"]}}'
            )
            resp = client_ai.chat.completions.create(
                model           = 'llama-3.1-8b-instant',
                messages        = [{
                    'role':    'user',
                    'content': prompt
                }],
                temperature     = 0.1,
                max_tokens      = 400,
                response_format = {
                    'type': 'json_object'}
            )
            raw    = resp.choices[0].message.content
            result = json.loads(raw)
            score  = max(0, min(
                int(result.get('score', 0)), marks))
            return {
                'score':      score,
                'out_of':     marks,
                'feedback':   str(
                    result.get('feedback', '')),
                'confidence': str(
                    result.get('confidence',
                               'medium')),
                'key_points': list(
                    result.get('key_points', [])),
                'model':      'groq/llama-3.1-8b-instant'
            }
        except Exception as e:
            print(f"[AI Grade] Error: {e}")

    # Keyword fallback
    stops = {
        'this','that','with','from','have','will',
        'what','when','where','which','they','their',
        'about','would','could','should','answer'
    }
    qw    = set(re.findall(
        r'\b[a-z]{4,}\b',
        question.lower())) - stops
    aw    = set(re.findall(
        r'\b[a-z]{4,}\b',
        str(answer).lower())) - stops
    wc    = len(str(answer).split())
    ratio = (
        min(1.0,
            (len(qw & aw) / max(len(qw), 1))
            + min(0.2, wc / 50))
        if qw
        else min(1.0, wc / 30)
    )
    score = round(marks * ratio)
    return {
        'score':      score,
        'out_of':     marks,
        'feedback':   f'Keyword score: {score}/{marks}.',
        'confidence': 'low',
        'key_points': [
            f'{len(qw & aw)}/{len(qw)} '
            f'key terms matched'],
        'model':      'keyword-fallback'
    }


@app.route(
    '/api/exams/<exam_id>/ai-grade/<user_id>',
    methods=['POST'])
@role_required('faculty')
def ai_grade_all(exam_id, user_id):
    try:
        rec = answers_col.find_one(
            {'exam_id': exam_id, 'user_id': user_id})
        if not rec:
            return jsonify(
                {'error': 'No submission'}), 404
        qs      = list(questions_col.find(
            {'exam_id': exam_id}
        ).sort('index', ASCENDING))
        answers = rec.get('answers', {})
        exam    = exams_col.find_one(
            {'exam_id': exam_id})
        subject = ((exam.get('subject') or 'General')
                   if exam else 'General')
        results  = []
        total_ai = 0
        skipped  = 0
        for q in qs:
            if q.get('type') != 'subjective':
                skipped += 1
                continue
            qid = str(q['_id'])
            g   = ai_grade(
                question = q.get('text', ''),
                answer   = str(
                    answers.get(qid, '')).strip(),
                marks    = int(q.get('marks', 1)),
                subject  = subject
            )
            total_ai += g['score']
            results.append({
                'question_id':    qid,
                'question_index': q.get('index', 0),
                'question_text':  q.get('text', ''),
                'student_answer': str(
                    answers.get(qid, '')),
                'marks':          int(
                    q.get('marks', 1)),
                **g
            })
        marks_col.update_one(
            {'exam_id': exam_id,
             'user_id': user_id},
            {'$set': {
                'exam_id':             exam_id,
                'user_id':             user_id,
                'ai_subjective_marks': total_ai,
                'ai_results':          results,
                'ai_graded_at':        utcnow()
            }},
            upsert=True
        )
        print(f"[AI Grade] ✅ {len(results)} graded, "
              f"{skipped} MCQ skipped, "
              f"total={total_ai}")
        return jsonify({
            'message':          (
                f'AI graded {len(results)} questions'),
            'total_ai_marks':   total_ai,
            'questions_graded': len(results),
            'results':          results
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route(
    '/api/exams/<exam_id>/ai-results/<user_id>',
    methods=['GET'])
@role_required('faculty')
def get_ai_results(exam_id, user_id):
    m = marks_col.find_one(
        {'exam_id': exam_id, 'user_id': user_id})
    if not m or 'ai_results' not in m:
        return jsonify({'graded': False}), 404
    graded_at = m.get('ai_graded_at', utcnow())
    return jsonify({
        'graded':              True,
        'ai_subjective_marks': m.get(
            'ai_subjective_marks', 0),
        'ai_graded_at':        (
            graded_at.isoformat()
            if hasattr(graded_at, 'isoformat')
            else str(graded_at)),
        'ai_model':            (
            m.get('ai_results') or [{}]
        )[0].get('model', 'unknown'),
        'results':             m.get('ai_results', [])
    })

# ═══════════════════════════════════════════════════════════════════════════
# ADMIN API
# ═══════════════════════════════════════════════════════════════════════════
def check_admin():
    auth = request.authorization
    return (auth
            and auth.username == ADMIN_USER
            and auth.password == ADMIN_PASS)


@app.route('/api/admin/users', methods=['GET'])
def admin_list_users():
    if not check_admin():
        return jsonify(
            {'error': 'Admin access only'}), 401
    role  = request.args.get('role', '')
    query = {'role': role} if role else {}
    users = list(users_col.find(
        query, {'password': 0}))
    for u in users:
        u['_id'] = str(u['_id'])
        for field in ['created_at', 'last_login']:
            if isinstance(u.get(field),
                          datetime.datetime):
                u[field] = u[field].isoformat()
    return jsonify(users)


@app.route(
    '/api/admin/users/<user_id>/toggle',
    methods=['POST'])
def admin_toggle(user_id):
    if not check_admin():
        return jsonify(
            {'error': 'Admin access only'}), 401
    oid  = get_oid(user_id)
    user = (users_col.find_one({'_id': oid})
            if oid else None)
    if not user:
        return jsonify(
            {'error': 'User not found'}), 404
    new = not user.get('is_active', True)
    users_col.update_one(
        {'_id': oid},
        {'$set': {'is_active': new}})
    return jsonify({'is_active': new})


@app.route('/api/admin/reset-password',
           methods=['POST'])
def admin_reset_password():
    if not check_admin():
        return jsonify(
            {'error': 'Admin access only'}), 401
    data     = request.get_json(silent=True) or {}
    roll     = data.get('roll_number', '').strip()
    new_pass = data.get('new_password', '').strip()
    if not roll or not new_pass:
        return jsonify({'error':
            'roll_number and new_password '
            'required'}), 400
    hashed = bcrypt.generate_password_hash(
        new_pass).decode('utf-8')
    r = users_col.update_one(
        {'roll_number': roll},
        {'$set': {'password': hashed}})
    if r.matched_count == 0:
        return jsonify(
            {'error': 'Student not found'}), 404
    return jsonify(
        {'message': f'Password reset for {roll}'})

# ═══════════════════════════════════════════════════════════════════════════
# SOCKET.IO
# ═══════════════════════════════════════════════════════════════════════════
@socketio.on('connect')
def on_connect():
    if 'user_id' in session:
        join_room(f"student_{session['user_id']}")

@socketio.on('join_exam')
def on_join(data):
    join_room(f"exam_{data.get('exam_id', '')}")
    emit('joined', {
        'room': f"exam_{data.get('exam_id', '')}"})

@socketio.on('leave_exam')
def on_leave(data):
    leave_room(f"exam_{data.get('exam_id', '')}")

@socketio.on('disconnect')
def on_disconnect():
    pass

# ═══════════════════════════════════════════════════════════════════════════
# STATIC IMAGES
# ═══════════════════════════════════════════════════════════════════════════
@app.route('/static/images/<path:filename>')
def static_images(filename):
    images_path = os.path.join(
        app.root_path, 'static', 'images')
    full_path   = os.path.join(images_path, filename)
    if os.path.exists(full_path):
        return send_from_directory(
            images_path, filename)
    return '', 404

# ═══════════════════════════════════════════════════════════════════════════
# STARTUP
# ═══════════════════════════════════════════════════════════════════════════
def print_banner():
    print("═" * 55)
    print("  ExamPro — Shobhit University Gangoh (SUG)")
    print("  Team Believer © 2026")
    print(f"  AI Parse   : llama3-70b-8192")
    print(f"  AI Grade   : llama-3.1-8b-instant")
    print(f"  Faculty Key: {FACULTY_REG_KEY}")
    print(f"  PDF Library: "
          f"{'fitz' if PYMUPDF_OK else 'pypdf' if PYPDF_OK else 'NONE'}")
    print(f"  Production : {IS_PRODUCTION}")
    print("═" * 55)


if __name__ == "__main__":
    print_banner()
    port  = int(os.environ.get('PORT', 5000))
    debug = not IS_PRODUCTION
    socketio.run(
        app,
        host  = '0.0.0.0',
        port  = port,
        debug = debug
    )
